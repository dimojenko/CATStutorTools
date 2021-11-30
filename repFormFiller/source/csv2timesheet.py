#!/usr/bin/python3
"""CSV to Timesheet Generator

This script reads in a (.csv) file of meetings and converts them to a MSWord
document (.docx) in CATS timesheet format. It is intended for use with CATS tutor
meetings, and the provided (.csv) file is expected to be formatted with the
following columns:
	Date, Student, Sport, Course, StartTime, EndTime

As long as the input file is formatted correctly, it may be included with or
without the header row.

This file can also be imported as a module. It is intended to be used in
conjunction with the calendar2csv.py module. 

command line usage:
	python3 csv2timesheet.py inputCSV -n [namesFile]
	- where inputCSV is the user provided (.csv) file
	- [namesFile] is an optional argument to provide a file (.txt) of names with
	  the format:
	  	tutor:
		tutor's last name, tutor's first name
		students:
		student's last name, student's first name
		...

The optional argument [namesFile] must be included for the first names of students
and the tutor's name to be included in the output document.
**Note: this option won't work properly if students share a last name

The output file will be named accordingly:
	If the namesFile is provided:
		If the input (.csv) is formatted as the output of calendar2csv.py:
			output file = [tutor's last name]_timesheet_[startDate]_to_[endDate].docx
		Else:
			output file = [tutor's last name]_timesheet.docx
	Else:
		If the input (.csv) is formatted as the output of calendar2csv.py:
			output file = timesheet_[startDate]_to_[endDate].docx
		Else:
			output file = timesheet.docx

requires:
	* the MSWord document "timesheetTemplate.docx" to be in runpath of csv2timesheet.py
"""

import sys
import argparse
import csv
from datetime import *
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from copy import deepcopy

def csv2timesheet(inputCSV, namesFile='none'):
	"""Given a (.csv) file of meetings, returns a (.docx) file in CATS timesheet format
	
	Parameters
	~~~~~~~~~~
	inputCSV : str
		The input (.csv) file of meetings

	namesFile : str, optional
		The input (.txt) file of tutor's and students' names in the format:
			tutor:
			[tutor's last name], [tutor's first name]
			students:
			[student's last name], [student's first name]
			...
	Returns
	~~~~~~~
	file(.docx)
		A MSWord document of the meetings in CATS timesheet format
	"""

	# open template document
	doc = Document('timesheetTemplate.docx')
	table = doc.tables[0]                        
	tableIndex = 0

	# if namesFile provided, assemble a list of names from it: [tutor, [student]]
	if not namesFile == 'none':
		if (checkFormat(namesFile)):
			namesList = namesFile2list(namesFile)
			tutor, students = namesList[0], namesList[1]
		else:
			print("The given namesFile doesn't have the correct format and won't be used.")
			namesFile = 'none'

	with open(inputCSV, 'r') as csvReader:
		inCSV = csv.reader(csvReader)
		tblRowCt = 0
		totalHours = 0
		totalSessions = 0
		for lnum, line in enumerate(inCSV):
			if not len(line) == 0 and not lnum == 0: # skip past header and blank rows
				newRow = table.add_row()
				tblRowCt += 1
				totalSessions += 1
				newRow.height = Inches(0.4)
				rowCells = newRow.cells
				for i in range(6):
					rowCells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
					if not i == 1:
						rowCells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # otherwise LEFT
					if i == 0: 
						# format date removing year
						dateSplit = line[i].split('/')
						date = dateSplit[0] + '/' + dateSplit[1]
						txtrun = rowCells[i].paragraphs[0].add_run(date)
						txtrun.bold = False # formatting is only applied to text when using runs
					elif i == 1:
						# find last name in namesFile if provided
						if not namesFile == 'none':
							fName = findFullName(line[i], students)
						else:
							fName = line[i]
						txtrun = rowCells[i].paragraphs[0].add_run(fName)
					else:
						txtrun = rowCells[i].paragraphs[0].add_run(line[i])

				# check for hours by difference between times
				sTime, eTime = line[4], line[5]
				sTimeSplit = sTime.split(':')
				sHrs, sMin = sTimeSplit[0], sTimeSplit[1]
				eTimeSplit = eTime.split(':')
				eHrs, eMin = eTimeSplit[0], eTimeSplit[1]
				sTime = datetime(2021, 1, 1, int(sHrs), int(sMin), 0) # assuming session times on same date
				eTime = datetime(2021, 1, 1, int(eHrs), int(eMin), 0)

				timeDiff = eTime - sTime # seconds stored in the timedelta object
				timeHours = timeDiff.total_seconds() / 3600 # convert to hours
				totalHours += timeHours
				rowCells[7].text = str(timeHours)
				rowCells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				rowCells[7].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

				# if 10 rows in table, create new table on new page
				if tblRowCt == 10:
					tblRowCt = 0
					newSec = doc.add_section()
					tabelCopy = deepcopy(table._tbl)
					p = doc.add_paragraph()
					p._p.addnext(tabelCopy)
					tableIndex += 1
					table = doc.tables[tableIndex]

					# remove rows from copy leaving the header
					for i in range(10):
						tbl = table._tbl
						row = table.rows[1]
						tr = row._tr
						tbl.remove(tr)

		p = doc.add_paragraph()
		seshText = "\nTotal Sessions: \t" + str(totalSessions)
		hrText = "\nTotal Hours:   \t" + str(totalHours)
		seshRun = p.add_run(seshText)
		seshRun.bold = True
		hrRun = p.add_run(hrText)
		hrRun.bold = True

	if 'meetings' in inputCSV:
		outFile = 'timesheet'+inputCSV.lstrip('meetings').rstrip('.csv')+'.docx'
	else:
		outFile = 'timesheet.docx'
	if not namesFile == 'none':
		lastName = tutor.split(',')[0].lower()
		outFile = lastName+'_'+outFile
	doc.save(outFile)
	print("Output file created: ", outFile)
	return outFile

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
# NOTICE: won't properly work for students that share a last name
def findFullName(lastName,namesList):
	"""Given a last name and a list of names, returns the full name"""
	fullName = lastName # return given name if not found in file
	for name in namesList:
		nameSplit = name.split(',')
		if nameSplit[0] == lastName:
			fullName = nameSplit[1]+' '+nameSplit[0]
	return fullName

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def namesFile2list(namesFile):
	"""Given a text file of names, returns a list: [tutor, [students]]"""
	namesList = []
	students = []
	with open(namesFile, 'r') as nfile:
		for line in nfile:
			if not len(line) == 0:
				line = line.strip()
				if line == 'tutor:':
					line = nfile.readline()
					namesList.append(line)
				elif line == 'students:':
					line = nfile.readline()
					while not len(line) == 0:
						students.append(line.strip())
						line = nfile.readline()
	namesList.append(students)
	return namesList	

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def checkFormat(namesFile):
	"""Checks the format of the file of names, returns True or False"""
	tutorLine, studentLine = False, False
	with open(namesFile, 'r') as nfile:
		for line in nfile:
			if not len(line) == 0:
				line = line.strip()
				if line == 'tutor:':
					tutorLine = True
				elif line == 'students:':
					studentLine = True
	return tutorLine and studentLine

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def main():
	argParser = argparse.ArgumentParser()
	argParser.add_argument(
		"inputCSV",
		type=str,
		help="The input (.csv) file of meetings with the format: Date,Student,Sport,Course,StartTime,EndTime"
	)
	argParser.add_argument(
		"-n", "--namesFile",
		type=str,
		default='none',
		help="""The input (.txt) file of tutor's name followed by students' names"""
	)
	args = argParser.parse_args()
	csv2timesheet(args.inputCSV, args.namesFile)

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
if __name__ == "__main__":
	main()